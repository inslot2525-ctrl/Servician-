import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document


def load_document(file_path: str):
    if file_path.endswith(".pdf"):
        return load_pdf(file_path)

    if file_path.endswith(".docx"):
        return load_docx(file_path)

    if file_path.endswith(".txt"):
        return load_txt(file_path)

    raise ValueError("Unsupported file format")


def load_pdf(file_path: str):
    docs = []
    pdf = fitz.open(file_path)

    for page_num, page in enumerate(pdf):
        text = page.get_text()
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path, "page": page_num + 1},
                )
            )

    return docs


def load_docx(file_path: str):
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    return [
        Document(
            page_content=text,
            metadata={"source": file_path, "page": 1},
        )
    ]


def load_txt(file_path: str):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="utf-16", errors="ignore") as f:
            text = f.read()

    return [
        Document(
            page_content=text,
            metadata={"source": file_path, "page": 1},
        )
    ]
